"""调度器测试"""

import json
import tempfile
from pathlib import Path

import pytest

from src.dispatcher import ParseDispatcher
from src.contract import ParseResult


@pytest.fixture
def dispatcher():
    """创建调度器实例"""
    return ParseDispatcher()


@pytest.fixture
def tmp_dir():
    """创建临时目录"""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


class TestDispatcher:
    """调度器核心测试"""

    def test_init(self, dispatcher):
        """初始化成功"""
        assert dispatcher is not None
        assert len(dispatcher.list_parsers()) > 0

    def test_list_supported_extensions(self, dispatcher):
        """列出支持的扩展名"""
        exts = dispatcher.list_supported_extensions()
        assert ".pdf" in exts
        assert ".docx" in exts
        assert ".md" in exts
        assert ".py" in exts
        assert ".jpg" in exts

    def test_can_parse(self, dispatcher):
        """判断能否解析"""
        assert dispatcher.can_parse(Path("test.pdf"))
        assert dispatcher.can_parse(Path("test.docx"))
        assert dispatcher.can_parse(Path("test.md"))
        assert dispatcher.can_parse(Path("test.py"))
        assert dispatcher.can_parse(Path("test.jpg"))
        # 不支持的格式回退到 text
        assert dispatcher.can_parse(Path("test.xyz")) is False

    def test_detect_type(self, dispatcher):
        """检测文件类型"""
        assert dispatcher.detect_type(Path("test.pdf")) == "pdf"
        assert dispatcher.detect_type(Path("test.docx")) == "office"
        assert dispatcher.detect_type(Path("test.md")) == "markdown"
        assert dispatcher.detect_type(Path("test.py")) == "code"
        assert dispatcher.detect_type(Path("test.jpg")) == "image"

    def test_list_parsers(self, dispatcher):
        """列出解析器"""
        parsers = dispatcher.list_parsers()
        names = [p["name"] for p in parsers]
        assert "pdf" in names
        assert "office" in names
        assert "image" in names
        assert "markdown" in names
        assert "text" in names
        assert "code" in names

    def test_dispatch_nonexistent_file(self, dispatcher):
        """不存在的文件"""
        result = dispatcher.dispatch(Path("/nonexistent/file.pdf"))
        assert result.success is False
        assert "不存在" in result.error

    def test_dispatch_markdown(self, dispatcher, tmp_dir):
        """解析 Markdown 文件"""
        md_file = tmp_dir / "test.md"
        md_file.write_text("# 标题\n\n这是测试内容。\n\n## 二级标题\n\n更多内容。", encoding="utf-8")

        result = dispatcher.dispatch(md_file)
        assert result.success is True
        assert result.type == "document"
        assert result.format == "md"
        assert "测试内容" in result.content
        assert result.metadata["heading_count"] == 2

    def test_dispatch_text(self, dispatcher, tmp_dir):
        """解析纯文本文件"""
        txt_file = tmp_dir / "test.txt"
        txt_file.write_text("Hello, World!\n这是测试。", encoding="utf-8")

        result = dispatcher.dispatch(txt_file)
        assert result.success is True
        assert result.type == "document"
        assert result.format == "txt"
        assert "Hello" in result.content
        assert "测试" in result.content

    def test_dispatch_docx(self, dispatcher, tmp_dir):
        """解析 Word 文档"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        docx_file = tmp_dir / "test.docx"
        doc = Document()
        doc.add_heading("测试标题", level=1)
        doc.add_paragraph("这是测试段落。")
        doc.save(str(docx_file))

        result = dispatcher.dispatch(docx_file)
        assert result.success is True
        assert result.type == "document"
        assert result.format == "docx"
        assert "测试标题" in result.content
        assert "测试段落" in result.content

    def test_dispatch_empty_file(self, dispatcher, tmp_dir):
        """空文件"""
        empty_file = tmp_dir / "empty.txt"
        empty_file.write_text("", encoding="utf-8")

        result = dispatcher.dispatch(empty_file)
        assert result.success is False

    def test_dispatch_code_python(self, dispatcher, tmp_dir):
        """解析 Python 文件"""
        py_file = tmp_dir / "test.py"
        py_file.write_text(
            '"""模块文档"""\n\n# 这是注释\ndef hello():\n    """函数文档"""\n    pass\n',
            encoding="utf-8",
        )

        result = dispatcher.dispatch(py_file)
        assert result.success is True
        assert result.type == "document"
        assert result.format == "py"
        assert "模块文档" in result.content or "注释" in result.content


class TestDispatcherContractCompliance:
    """验证所有解析器输出符合契约"""

    def _check_contract(self, result: ParseResult):
        """检查结果符合契约"""
        assert result.source, "source 不能为空"
        assert result.type in ("document", "audio", "image", "video"), f"无效 type: {result.type}"
        assert result.format, "format 不能为空"
        assert isinstance(result.content, str), "content 必须是字符串"
        assert isinstance(result.metadata, dict), "metadata 必须是字典"

    def test_markdown_contract(self, dispatcher, tmp_dir):
        """Markdown 解析器输出符合契约"""
        f = tmp_dir / "test.md"
        f.write_text("# Hello\n\nWorld", encoding="utf-8")
        result = dispatcher.dispatch(f)
        self._check_contract(result)

    def test_text_contract(self, dispatcher, tmp_dir):
        """文本解析器输出符合契约"""
        f = tmp_dir / "test.txt"
        f.write_text("Hello World", encoding="utf-8")
        result = dispatcher.dispatch(f)
        self._check_contract(result)

    def test_docx_contract(self, dispatcher, tmp_dir):
        """Word 解析器输出符合契约"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        f = tmp_dir / "test.docx"
        doc = Document()
        doc.add_paragraph("Hello")
        doc.save(str(f))
        result = dispatcher.dispatch(f)
        self._check_contract(result)
